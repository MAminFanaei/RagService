# test_quality_docx.py

import json
import time
from pathlib import Path

def save_result(output_path, method, data):
    with open(output_path / f"{method}.json", "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"  ✅ {method:<20} → {output_path / f'{method}.json'}")

def run_docx_tests(file_path: Path, output_dir: Path, args):
    print(f"📄 Testing DOCX parser on: {file_path.name}\n")
    
    print("▶ Testing python-docx...")
    try:
        from docx import Document
        from tabulate import tabulate
        
        t0 = time.perf_counter()
        doc = Document(file_path)
        
        # Extract metadata
        props = doc.core_properties
        metadata = {
            "title": props.title,
            "author": props.author,
            "created": str(props.created),
            "modified": str(props.modified),
        }
        
        # Extract structure
        paragraphs = []
        headings = []
        tables_data = []
        
        for para in doc.paragraphs:
            style = para.style.name
            text = para.text.strip()
            
            if not text:
                continue
            
            if style.startswith("Heading"):
                level = int(style.split()[-1]) if style != "Heading" else 1
                headings.append({
                    "level": level,
                    "text": text,
                })
            
            paragraphs.append({
                "style": style,
                "text": text[:500],
                "is_heading": style.startswith("Heading"),
            })
        
        # Extract tables
        for table in doc.tables:
            rows_data = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])
            
            # Convert to markdown
            if rows_data:
                markdown = tabulate(rows_data[1:], headers=rows_data[0], tablefmt="github")
                tables_data.append({
                    "row_count": len(rows_data),
                    "col_count": len(rows_data[0]) if rows_data else 0,
                    "markdown": markdown[:1000],
                })
        
        elapsed = time.perf_counter() - t0
        
        result = {
            "parser": "python-docx",
            "file": file_path.name,
            "time_seconds": round(elapsed, 3),
            "metadata": metadata,
            "paragraph_count": len(paragraphs),
            "heading_count": len(headings),
            "table_count": len(tables_data),
            "headings_preview": headings[:20],
            "paragraphs_preview": paragraphs[:20],
            "tables_preview": tables_data[:5],
            "verdict": {
                "has_structure": len(headings) > 0,
                "has_tables": len(tables_data) > 0,
                "quality": "✅ Rich structure" if headings and tables_data else "⚠️ Plain text only"
            }
        }
        save_result(output_dir, "parser_python_docx", result)
        
        # Now chunk the full text
        full_text = "\n\n".join(p["text"] for p in paragraphs)
        
        print(f"\n📦 Testing chunkers on extracted text...\n")
        
        # Fixed chunker
        print("▶ Testing Fixed chunker...")
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            
            t0 = time.perf_counter()
            splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
            chunks = splitter.split_text(full_text)
            elapsed = time.perf_counter() - t0
            
            save_result(output_dir, "chunker_fixed", {
                "chunker": "fixed",
                "chunk_count": len(chunks),
                "time_seconds": round(elapsed, 3),
                "chunks_preview": [{"index": i, "text": c[:500]} for i, c in enumerate(chunks[:args.max_chunks])],
            })
        except Exception as e:
            print(f"  ❌ Failed: {e}")
        
    except ImportError:
        print(f"  ⏭️  python-docx not installed — skip")
    except Exception as e:
        print(f"  ❌ python-docx failed: {e}")