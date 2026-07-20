"""
DOCX parser using python-docx.

Extracts:
  - Paragraphs with heading-level detection → section_path breadcrumb
  - Tables → Markdown via tabulate
  - Footnotes → via XML namespace w:footnote
  - Document properties (title, author, created, modified)

Fixes applied from old pipeline (as described in the plan):
  - h7/h8 were overwriting h6 key → fixed with correct dict keys
  - duplicate `elif lvl == 6` → fixed with lvl 7 & 8
  - grouping by (doc_id, section_id) where doc_id was None → removed

Note: DOCX does not store page numbers so page_number is always None.
"""

from __future__ import annotations

import io
import re
from typing import Any

import structlog

from ingestion.parsers.base import BaseParser, ParsedElement, ParserError

logger = structlog.get_logger(__name__)

# ---------------------------------------------------------------------------
# Heading level extraction
# ---------------------------------------------------------------------------

_HEADING_STYLE_RE = re.compile(r"^[Hh]eading\s*(\d+)$")


def _heading_level(style_name: str) -> int | None:
    """Return heading level (1-8) for a Word paragraph style, or None."""
    m = _HEADING_STYLE_RE.match(style_name.strip())
    if m:
        return int(m.group(1))
    return None


# ---------------------------------------------------------------------------
# Section path management
# ---------------------------------------------------------------------------

class _SectionTracker:
    """Maintains a breadcrumb stack of heading texts up to heading level 8."""

    def __init__(self) -> None:
        # Keys 1-8, values are the most-recent heading text at that level
        self._headings: dict[int, str] = {}

    def update(self, level: int, text: str) -> list[str]:
        """
        Record a new heading at `level` and return the updated breadcrumb.
        Headings below `level` are cleared (new parent resets children).
        """
        # Clear all deeper levels
        for lvl in range(level + 1, 9):
            self._headings.pop(lvl, None)
        self._headings[level] = text
        return self._current_path()

    def current(self) -> list[str]:
        return self._current_path()

    def current_title(self) -> str | None:
        if not self._headings:
            return None
        max_lvl = max(self._headings)
        return self._headings[max_lvl]

    def _current_path(self) -> list[str]:
        return [self._headings[lvl] for lvl in sorted(self._headings)]


# ---------------------------------------------------------------------------
# Table → Markdown
# ---------------------------------------------------------------------------

def _table_to_markdown(table) -> str:
    """Convert a python-docx Table to a Markdown string."""
    try:
        from tabulate import tabulate  # type: ignore

        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        if not rows:
            return ""

        # First row as header if it looks like headers
        return tabulate(rows[1:], headers=rows[0], tablefmt="pipe") if len(rows) > 1 else tabulate(rows, tablefmt="pipe")
    except ImportError:
        # tabulate not installed — fall back to simple pipe table
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if not rows:
            return ""
        lines = [" | ".join(rows[0]), " | ".join(["---"] * len(rows[0]))]
        lines += [" | ".join(r) for r in rows[1:]]
        return "\n".join(lines)


# ---------------------------------------------------------------------------
# Footnote extraction
# ---------------------------------------------------------------------------

def _extract_footnotes(docx_obj) -> list[str]:
    """
    Extract footnote texts via python-docx's underlying lxml element tree.
    Uses the standard DOCX XML namespace for footnotes.
    """
    footnotes: list[str] = []
    try:
        from lxml import etree  # type: ignore

        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        part = docx_obj.part

        # Footnotes part may not exist in all documents
        try:
            footnotes_part = part.footnotes_part
        except AttributeError:
            return []

        root = footnotes_part._element
        for fn in root.findall(f"{{{ns}}}footnote"):
            fn_id = fn.get(f"{{{ns}}}id", "")
            # Skip separator pseudo-footnotes (id -1, 0)
            if fn_id in ("-1", "0"):
                continue
            texts = []
            for para in fn.findall(f".//{{{ns}}}p"):
                for run in para.findall(f".//{{{ns}}}r"):
                    for t in run.findall(f"{{{ns}}}t"):
                        if t.text:
                            texts.append(t.text)
            combined = "".join(texts).strip()
            if combined:
                footnotes.append(combined)
    except Exception as exc:
        logger.warning("docx_footnote_extraction_failed", error=str(exc))

    return footnotes


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

class DocxParser(BaseParser):
    """Primary parser for .docx files."""

    @property
    def name(self) -> str:
        return "docx"

    async def parse(self, data: bytes, filename: str) -> list[ParsedElement]:
        try:
            import docx as python_docx  # python-docx
        except ImportError as exc:
            raise ParserError(self.name, "python-docx not installed", exc)

        try:
            doc = python_docx.Document(io.BytesIO(data))
        except Exception as exc:
            raise ParserError(self.name, f"Failed to open DOCX: {exc}", exc)

        elements: list[ParsedElement] = []
        tracker = _SectionTracker()

        # --- Document properties ---
        props: dict[str, Any] = {}
        try:
            cp = doc.core_properties
            props = {
                "title": cp.title or "",
                "author": cp.author or "",
                "created": str(cp.created) if cp.created else "",
                "modified": str(cp.modified) if cp.modified else "",
            }
        except Exception:
            pass

        # --- Iterate body elements in document order ---
        # python-docx exposes doc.paragraphs and doc.tables separately; to
        # preserve order we walk doc.element.body children.
        try:
            from docx.oxml.ns import qn  # type: ignore
        except ImportError:
            qn = None  # type: ignore

        body = doc.element.body

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            # ---- Paragraph ----
            if tag == "p":
                para = self._wrap_paragraph(child, doc)
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue

                lvl = _heading_level(para.style.name)

                if lvl is not None:
                    # It's a heading
                    path = tracker.update(lvl, text)
                    elements.append(
                        ParsedElement(
                            text=text,
                            parser_name=self.name,
                            element_type="heading",
                            heading_level=lvl,
                            section_path=list(path),
                            section_title=text,
                            raw_metadata={**props},
                        )
                    )
                else:
                    # Regular paragraph
                    style_lower = para.style.name.lower()
                    if "footnote" in style_lower:
                        etype = "footnote"
                    elif "caption" in style_lower:
                        etype = "caption"
                    elif "code" in style_lower or "verbatim" in style_lower:
                        etype = "code"
                    else:
                        etype = "text"

                    elements.append(
                        ParsedElement(
                            text=text,
                            parser_name=self.name,
                            element_type=etype,
                            section_title=tracker.current_title(),
                            section_path=tracker.current(),
                            raw_metadata={**props},
                        )
                    )

            # ---- Table ----
            elif tag == "tbl":
                tbl = self._wrap_table(child, doc)
                if tbl is None:
                    continue
                md = _table_to_markdown(tbl)
                plain = "\n".join(
                    " | ".join(cell.text.strip() for cell in row.cells)
                    for row in tbl.rows
                )
                if not plain.strip():
                    continue
                elements.append(
                    ParsedElement(
                        text=plain,
                        parser_name=self.name,
                        element_type="table",
                        is_table=True,
                        table_markdown=md,
                        section_title=tracker.current_title(),
                        section_path=tracker.current(),
                        raw_metadata={**props},
                    )
                )

        # --- Footnotes ---
        footnotes = _extract_footnotes(doc)
        for fn_text in footnotes:
            elements.append(
                ParsedElement(
                    text=fn_text,
                    parser_name=self.name,
                    element_type="footnote",
                    section_title=tracker.current_title(),
                    section_path=tracker.current(),
                    raw_metadata={**props},
                )
            )

        logger.info(
            "docx_parsed",
            filename=filename,
            element_count=len(elements),
        )
        return elements

    # ------------------------------------------------------------------
    # Helpers to wrap raw lxml elements back into python-docx objects
    # ------------------------------------------------------------------

    def _wrap_paragraph(self, element, doc):
        """Wrap a raw lxml `<w:p>` element as a python-docx Paragraph."""
        try:
            from docx.text.paragraph import Paragraph  # type: ignore
            return Paragraph(element, doc)
        except Exception:
            return None

    def _wrap_table(self, element, doc):
        """Wrap a raw lxml `<w:tbl>` element as a python-docx Table."""
        try:
            from docx.table import Table  # type: ignore
            return Table(element, doc)
        except Exception:
            return None
